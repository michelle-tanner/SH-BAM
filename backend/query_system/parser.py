"""
parser.py
---------
Extract plain text from PDF / DOCX / TXT and infer a document date from the
filename or file modification time (ISO YYYY-MM-DD).
"""

from __future__ import annotations

import re
from datetime import datetime, timezone
from pathlib import Path


_DATE_IN_NAME = re.compile(r"(20\d{2})[-_]?(\d{2})[-_]?(\d{2})|(\d{4})(\d{2})(\d{2})")


def guess_doc_date(path: Path) -> str:
    """Best-effort YYYY-MM-DD from filename patterns, else file mtime."""
    stem = path.stem
    m = _DATE_IN_NAME.search(stem)
    if m:
        if m.group(1):
            y, mo, d = m.group(1), m.group(2), m.group(3)
        else:
            y, mo, d = m.group(4), m.group(5), m.group(6)
        try:
            datetime(int(y), int(mo), int(d))
            return f"{y}-{mo}-{d}"
        except ValueError:
            pass
    try:
        ts = path.stat().st_mtime
        return datetime.fromtimestamp(ts, tz=timezone.utc).strftime("%Y-%m-%d")
    except OSError:
        return datetime.now(timezone.utc).strftime("%Y-%m-%d")


def parse_document(path: Path) -> str:
    """Return extracted UTF-8 text for supported file types."""
    path = path.resolve()
    suffix = path.suffix.lower()

    if suffix == ".txt":
        return path.read_text(encoding="utf-8", errors="replace")

    if suffix == ".pdf":
        try:
            import fitz  # PyMuPDF

            doc = fitz.open(path)
            try:
                text = "\n\n".join(page.get_text() for page in doc)
            finally:
                doc.close()
            if text.strip():
                return text
            # Empty → image-based/scanned PDF; fall through to unstructured
        except ImportError:
            pass

    if suffix == ".docx":
        try:
            from docx import Document

            doc = Document(str(path))
            parts: list[str] = []
            # Paragraphs (body text, headings)
            parts.extend(p.text for p in doc.paragraphs if p.text.strip())
            # Tables — python-docx only reads paragraphs by default; tables are skipped
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            parts.append(cell.text.strip())
            text = "\n\n".join(parts)
            if text.strip():
                return text
            # Empty → content may be in text boxes or complex layout; fall through
        except ImportError:
            pass

    # Last resort: unstructured handles edge cases (image PDFs with OCR, complex DOCX, etc.)
    from unstructured.partition.auto import partition

    elements = partition(filename=str(path))
    return "\n\n".join(str(el).strip() for el in elements if str(el).strip())
